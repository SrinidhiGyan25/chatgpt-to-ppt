#!/usr/bin/env python3
"""
Document to JSON Mapping Generator
Reads Word documents and creates JSON mapping files for PowerPoint image insertion
with collision detection and automatic position assignment
"""

import os
import sys
import json
import re
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple
from dataclasses import dataclass
from enum import Enum
import docx
from docx import Document

class Position(Enum):
    TOP_LEFT = "top-left"
    TOP_RIGHT = "top-right"
    BOTTOM_LEFT = "bottom-left"
    BOTTOM_RIGHT = "bottom-right"
    CENTER = "center"
    CUSTOM = "custom"

@dataclass
class SlideMapping:
    slide_number: int
    image_numbers: List[int]
    positions: List[Position]

class PositionManager:
    """Manages position assignment with collision detection"""
    
    def __init__(self):
        self.available_positions = [
            Position.BOTTOM_LEFT,
            Position.BOTTOM_RIGHT,
            Position.TOP_RIGHT
        ]
        # Track occupied positions per slide
        self.occupied_positions: Dict[int, Set[Position]] = {}
    
    def get_next_available_position(self, slide_number: int) -> Optional[Position]:
        """Get next available position on a slide"""
        if slide_number not in self.occupied_positions:
            self.occupied_positions[slide_number] = set()
        
        for position in self.available_positions:
            if position not in self.occupied_positions[slide_number]:
                return position
        return None
    
    def occupy_position(self, slide_number: int, position: Position):
        """Mark position as occupied"""
        if slide_number not in self.occupied_positions:
            self.occupied_positions[slide_number] = set()
        self.occupied_positions[slide_number].add(position)
    
    def is_slide_full(self, slide_number: int) -> bool:
        """Check if all positions on slide are occupied"""
        return self.get_next_available_position(slide_number) is None
    
    def get_available_count(self, slide_number: int) -> int:
        """Get count of available positions on slide"""
        if slide_number not in self.occupied_positions:
            return len(self.available_positions)
        return len(self.available_positions) - len(self.occupied_positions[slide_number])

class DocumentParser:
    """Parse Word document for slide-image mappings"""
    
    def __init__(self):
        self.slide_patterns = [
            r'slide\s*:?\s*(\d+)',  # "slide: 1" or "slide 1"
            r'slide\s*#\s*(\d+)',   # "slide #1"
            r'slide\s*number\s*:?\s*(\d+)',  # "slide number: 1"
            r'page\s*:?\s*(\d+)',   # "page: 1"
            r'pg\s*:?\s*(\d+)',     # "pg: 1"
        ]
        
        self.image_patterns = [
            r'image\s*:?\s*(\d+(?:\s*,\s*\d+)*)',  # "image: 1,2,3"
            r'img\s*:?\s*(\d+(?:\s*,\s*\d+)*)',    # "img: 1,2,3"
            r'picture\s*:?\s*(\d+(?:\s*,\s*\d+)*)', # "picture: 1,2,3"
            r'photo\s*:?\s*(\d+(?:\s*,\s*\d+)*)',   # "photo: 1,2,3"
            r'images?\s*:?\s*(\d+(?:\s*[-,]\s*\d+)*)', # "images: 1-3" or "image: 1,2,3"
        ]
    
    def parse_document(self, doc_path: str) -> List[SlideMapping]:
        """Parse Word document and extract slide mappings"""
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"Document '{doc_path}' not found")
        
        try:
            doc = Document(doc_path)
            print(f"✓ Loaded document: {doc_path}")
        except Exception as e:
            raise Exception(f"Error loading document: {e}")
        
        # Extract text from all paragraphs and tables
        full_text = self._extract_all_text(doc)
        
        # Parse slide mappings
        mappings = self._parse_slide_mappings(full_text)
        
        return mappings
    
    def _extract_all_text(self, doc: Document) -> str:
        """Extract all text from document including tables"""
        text_parts = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        return '\n'.join(text_parts)
    
    def _parse_slide_mappings(self, text: str) -> List[SlideMapping]:
        """Parse slide mappings from text"""
        mappings = []
        lines = text.split('\n')
        
        current_slide = None
        
        for line_num, line in enumerate(lines, 1):
            line = line.strip().lower()
            if not line:
                continue
            
            # Check for slide number
            slide_match = self._find_slide_number(line)
            if slide_match:
                current_slide = int(slide_match)
                print(f"  Found slide: {current_slide} (line {line_num})")
                continue
            
            # Check for image numbers
            if current_slide is not None:
                image_numbers = self._find_image_numbers(line)
                if image_numbers:
                    print(f"  Found images for slide {current_slide}: {image_numbers}")
                    mappings.append(SlideMapping(
                        slide_number=current_slide,
                        image_numbers=image_numbers,
                        positions=[]  # Will be assigned later
                    ))
                    current_slide = None  # Reset after finding images
        
        return mappings
    
    def _find_slide_number(self, line: str) -> Optional[int]:
        """Find slide number in line"""
        for pattern in self.slide_patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                return int(match.group(1))
        return None
    
    def _find_image_numbers(self, line: str) -> List[int]:
        """Find image numbers in line"""
        for pattern in self.image_patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                numbers_str = match.group(1)
                return self._parse_number_list(numbers_str)
        return []
    
    def _parse_number_list(self, numbers_str: str) -> List[int]:
        """Parse number list from string (handles ranges and comma-separated)"""
        numbers = []
        
        # Handle ranges (1-3) and comma-separated (1,2,3)
        parts = re.split(r'[,\s]+', numbers_str.strip())
        
        for part in parts:
            if '-' in part and not part.startswith('-'):
                # Handle range (e.g., "1-3")
                try:
                    start, end = map(int, part.split('-', 1))
                    numbers.extend(range(start, end + 1))
                except ValueError:
                    # If range parsing fails, try as single number
                    try:
                        numbers.append(int(part))
                    except ValueError:
                        continue
            else:
                # Single number
                try:
                    numbers.append(int(part))
                except ValueError:
                    continue
        
        return sorted(list(set(numbers)))  # Remove duplicates and sort

class JSONGenerator:
    """Generate JSON mapping file with position assignment"""
    
    def __init__(self):
        self.position_manager = PositionManager()
    
    def generate_json_mapping(self, mappings: List[SlideMapping], 
                            output_file: str = "mapping.json",
                            default_width: float = 3.0,
                            default_height: float = 2.0) -> bool:
        """Generate JSON mapping file with collision detection"""
        
        print(f"\n=== Generating JSON Mapping ===")
        print(f"Processing {len(mappings)} slide mappings...")
        
        # First pass: analyze all mappings to detect potential collisions
        self._analyze_slide_usage(mappings)
        
        # Second pass: assign positions with collision detection
        json_entries = []
        
        for mapping in mappings:
            slide_num = mapping.slide_number
            image_numbers = mapping.image_numbers
            
            print(f"\nProcessing slide {slide_num} with {len(image_numbers)} images:")
            
            # Check if slide can accommodate all images
            available_positions = self.position_manager.get_available_count(slide_num)
            
            if len(image_numbers) > available_positions:
                print(f"  ⚠ Warning: Slide {slide_num} needs {len(image_numbers)} positions but only has {available_positions} available")
                print(f"  ✓ Will distribute excess images to new slides")
            
            # Assign positions for each image
            for i, img_num in enumerate(image_numbers):
                # Check if current slide has space
                if self.position_manager.is_slide_full(slide_num):
                    # Find next available slide or create new one
                    slide_num = self._find_or_create_slide(slide_num)
                    print(f"  ✓ Moving to slide {slide_num} for image {img_num}")
                
                # Get next available position
                position = self.position_manager.get_next_available_position(slide_num)
                if position:
                    self.position_manager.occupy_position(slide_num, position)
                    
                    entry = {
                        "image_number": img_num,
                        "slide_number": slide_num,
                        "position": position.value,
                        "width": default_width,
                        "height": default_height
                    }
                    
                    json_entries.append(entry)
                    print(f"    Image {img_num} → Slide {slide_num}, Position: {position.value}")
                else:
                    print(f"  ✗ Error: Could not assign position for image {img_num}")
        
        # Write JSON file
        try:
            with open(output_file, 'w') as f:
                json.dump(json_entries, f, indent=2)
            
            print(f"\n✓ Successfully created JSON mapping: {output_file}")
            print(f"✓ Generated {len(json_entries)} image placements")
            
            # Print summary
            self._print_assignment_summary(json_entries)
            return True
            
        except Exception as e:
            print(f"\n✗ Error writing JSON file: {e}")
            return False
    
    def _analyze_slide_usage(self, mappings: List[SlideMapping]):
        """Analyze slide usage to identify potential issues"""
        slide_usage = {}
        for mapping in mappings:
            slide_num = mapping.slide_number
            image_count = len(mapping.image_numbers)
            
            if slide_num in slide_usage:
                slide_usage[slide_num] += image_count
            else:
                slide_usage[slide_num] = image_count
        
        print("\n=== Slide Usage Analysis ===")
        max_positions = len(self.position_manager.available_positions)
        
        for slide_num, count in sorted(slide_usage.items()):
            if count > max_positions:
                print(f"Slide {slide_num}: {count} images (exceeds {max_positions} positions)")
            else:
                print(f"Slide {slide_num}: {count} images")
    
    def _find_or_create_slide(self, current_slide: int) -> int:
        """Find next available slide or suggest new slide number"""
        # Start from current slide + 1 and find next available
        next_slide = current_slide + 1
        
        # For now, we'll just increment slide numbers
        # In a real implementation, you might want to check existing presentation
        while self.position_manager.is_slide_full(next_slide):
            next_slide += 1
        
        return next_slide
    
    def _print_assignment_summary(self, json_entries: List[dict]):
        """Print summary of position assignments"""
        print("\n=== Assignment Summary ===")
        
        # Group by slide
        slides = {}
        for entry in json_entries:
            slide_num = entry["slide_number"]
            if slide_num not in slides:
                slides[slide_num] = []
            slides[slide_num].append(entry)
        
        for slide_num in sorted(slides.keys()):
            entries = slides[slide_num]
            positions = [entry["position"] for entry in entries]
            images = [entry["image_number"] for entry in entries]
            
            print(f"Slide {slide_num}: Images {images} → Positions {positions}")

def create_sample_document():
    """Create a sample Word document for testing"""
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading('Sample Slide-Image Mapping Document', 0)
        
        # Add some sample mappings
        doc.add_paragraph('This document contains slide and image mappings for PowerPoint presentation.')
        doc.add_paragraph('')
        
        # Sample mappings
        mappings = [
            ("Slide: 1", "Images: 1, 2"),
            ("Slide: 2", "Image: 3"),
            ("Slide: 3", "Images: 4-6"),
            ("Slide: 4", "Pictures: 7, 8, 9, 10"),  # This will exceed available positions
            ("Slide: 5", "Photo: 11"),
        ]
        
        for slide_info, image_info in mappings:
            doc.add_paragraph(slide_info, style='Heading 2')
            doc.add_paragraph(image_info)
            doc.add_paragraph('')  # Empty line
        
        # Add table example
        doc.add_heading('Alternative Table Format', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Slide Number'
        hdr_cells[1].text = 'Image Numbers'
        
        # Add data rows
        table_data = [
            ('6', '12, 13'),
            ('7', '14-16'),
            ('8', '17'),
        ]
        
        for slide, images in table_data:
            row_cells = table.add_row().cells
            row_cells[0].text = f"Slide: {slide}"
            row_cells[1].text = f"Images: {images}"
        
        # Save document
        doc.save('sample_mapping.docx')
        print("✓ Created sample Word document: sample_mapping.docx")
        return True
        
    except Exception as e:
        print(f"✗ Error creating sample document: {e}")
        return False

def main():
    if len(sys.argv) > 1 and sys.argv[1] == '--create-sample':
        create_sample_document()
        return
    
    print("=== Document to JSON Mapping Generator ===\n")
    
    try:
        # Get inputs
        doc_path = input("Enter Word document path (.docx): ").strip()
        if not doc_path:
            print("Document path is required!")
            return
        
        output_file = input("Enter output JSON file name (press Enter for 'mapping.json'): ").strip()
        if not output_file:
            output_file = "mapping.json"
        
        # Get default dimensions
        try:
            width = input("Enter default image width in inches (press Enter for 3.0): ").strip()
            width = float(width) if width else 3.0
            
            height = input("Enter default image height in inches (press Enter for 2.0): ").strip()
            height = float(height) if height else 2.0
        except ValueError:
            print("Invalid dimensions, using defaults (3.0 x 2.0)")
            width, height = 3.0, 2.0
        
        # Parse document
        parser = DocumentParser()
        mappings = parser.parse_document(doc_path)
        
        if not mappings:
            print("\n⚠ No slide-image mappings found in document!")
            print("Make sure your document contains patterns like:")
            print("  - 'Slide: 1' followed by 'Images: 1,2,3'")
            print("  - 'Page: 2' followed by 'Pictures: 4-6'")
            return
        
        # Generate JSON
        generator = JSONGenerator()
        success = generator.generate_json_mapping(mappings, output_file, width, height)
        
        if success:
            print(f"\n✓ Process completed successfully!")
            print(f"✓ You can now use '{output_file}' with the PowerPoint image inserter")
        else:
            print("\n✗ Process failed!")
        
    except KeyboardInterrupt:
        print("\n\nOperation cancelled by user.")
    except Exception as e:
        print(f"\nUnexpected error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()